from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"


def set_run_font(run, size: float | None = None, bold: bool = False, color: str | None = None, font: str = "Arial") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.bold = bold
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_paragraph(paragraph, color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    p_pr.append(shd)


def set_paragraph_spacing(paragraph, before: float = 0, after: float = 8, line: float = 1.15) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = line


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color in [
        ("Heading 1", 20, "000000"),
        ("Heading 2", 16, "000000"),
        ("Heading 3", 13, "434343"),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(14 if style_name != "Heading 1" else 18)
        style.paragraph_format.space_after = Pt(6)


def add_inline_runs(paragraph, text: str, size: float = 11) -> None:
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=size, font="Consolas")
            run.font.color.rgb = RGBColor.from_string("374151")
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size)


def add_code_block(doc: Document, lines: list[str]) -> None:
    if not lines:
        return
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=4, after=8, line=1.05)
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    shade_paragraph(paragraph, "F3F4F6")
    run = paragraph.add_run("\n".join(lines))
    set_run_font(run, size=9.5, font="Consolas")
    run.font.color.rgb = RGBColor.from_string("111827")


def build_docx(markdown_path: Path, output_path: Path) -> None:
    doc = Document()
    configure_document(doc)

    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_lines: list[str] = []

    for raw in lines:
        line = raw.rstrip()

        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not line.strip():
            continue

        if line.startswith("# "):
            paragraph = doc.add_paragraph()
            set_paragraph_spacing(paragraph, before=0, after=10, line=1.15)
            run = paragraph.add_run(line[2:].strip())
            set_run_font(run, size=24, bold=False)
            continue

        if line.startswith("## "):
            paragraph = doc.add_paragraph(line[3:].strip(), style="Heading 1")
            set_paragraph_spacing(paragraph, before=18, after=6, line=1.15)
            continue

        if line.startswith("### "):
            paragraph = doc.add_paragraph(line[4:].strip(), style="Heading 2")
            set_paragraph_spacing(paragraph, before=14, after=5, line=1.15)
            continue

        bullet_match = re.match(r"^- (.+)", line)
        if bullet_match:
            paragraph = doc.add_paragraph(style="List Bullet")
            set_paragraph_spacing(paragraph, before=0, after=4, line=1.15)
            add_inline_runs(paragraph, bullet_match.group(1))
            continue

        number_match = re.match(r"^\d+\. (.+)", line)
        if number_match:
            paragraph = doc.add_paragraph(style="List Number")
            set_paragraph_spacing(paragraph, before=0, after=4, line=1.15)
            add_inline_runs(paragraph, number_match.group(1))
            continue

        paragraph = doc.add_paragraph()
        set_paragraph_spacing(paragraph, before=0, after=8, line=1.15)
        add_inline_runs(paragraph, line)

    if in_code:
        add_code_block(doc, code_lines)

    doc.core_properties.title = markdown_path.stem
    doc.core_properties.author = "VisionCraft"
    doc.save(output_path)


def main() -> None:
    targets = [
        DOCS_DIR / "代码说明文档.md",
        DOCS_DIR / "使用文档与HCI分析.md",
    ]
    for markdown_path in targets:
        output_path = markdown_path.with_suffix(".docx")
        build_docx(markdown_path, output_path)
        print(output_path)


if __name__ == "__main__":
    main()
